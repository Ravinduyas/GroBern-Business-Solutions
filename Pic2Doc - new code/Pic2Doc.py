import os
from tkinter import *
from tkinter import filedialog, messagebox
from PIL import Image, ImageEnhance
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import re
from tkinter import ttk

# Function to select folder path using tkinter file dialog
def select_folder():
    folder_path = filedialog.askdirectory()
    if folder_path:
        progress_bar['value'] = 0  # Reset progress bar
        process_images(folder_path)
    else:
        messagebox.showwarning("Warning", "Please select a folder.")

# Function to update progress bar
def update_progress_bar(current_value, total):
    progress = int((current_value / total) * 100)
    progress_bar['value'] = progress
    root.update_idletasks()  # Refresh the GUI

# Function to process images in the selected folder
def process_images(input_folder):
    try:
        # Create an output folder for enhanced images
        output_folder = os.path.join(input_folder, 'enhanced')
        os.makedirs(output_folder, exist_ok=True)

        # Enhancement parameters
        brightness_factor = 1.5  # Adjust as needed
        contrast_factor = 2.5    # Adjust as needed
        sharpness_factor = 5.0   # Adjust as needed

        # Create a new Word document
        doc = Document()

        # Set page layout to A4 and margins to narrow
        sections = doc.sections
        for section in sections:
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Custom sort function to handle the specific filename format
        def sort_images(filenames):
            def extract_timestamp_and_index(filename):
                match = re.search(r'WhatsApp Image (\d{4}-\d{2}-\d{2} at \d{2}\.\d{2}\.\d{2})(?: \((\d+)\))?', filename)
                if match:
                    timestamp_str = match.group(1)
                    index = int(match.group(2)) if match.group(2) else 0
                    timestamp = timestamp_str.replace(' at ', ' ')
                    return timestamp, index
                else:
                    return '', float('inf')

            return sorted(filenames, key=extract_timestamp_and_index, reverse=True)

        # Get sorted list of images
        image_files = sort_images([f for f in os.listdir(input_folder) if f.lower().endswith(('png', 'jpg', 'jpeg', 'bmp', 'gif', 'tiff'))])

        total_images = len(image_files)

        # Process each image in the folder
        for index, filename in enumerate(image_files, start=1):
            # Open an image file
            img_path = os.path.join(input_folder, filename)
            with Image.open(img_path) as img:
                # Convert image to RGB mode if not already
                if img.mode != 'RGB':
                    img = img.convert('RGB')

                # Enhance brightness
                enhancer = ImageEnhance.Brightness(img)
                img = enhancer.enhance(brightness_factor)

                # Enhance contrast
                enhancer = ImageEnhance.Contrast(img)
                img = enhancer.enhance(contrast_factor)

                # Enhance sharpness
                enhancer = ImageEnhance.Sharpness(img)
                img = enhancer.enhance(sharpness_factor)

                # Save the enhanced image to the output folder
                output_path = os.path.join(output_folder, filename)
                try:
                    img.save(output_path)
                except Exception as e:
                    print(f"Error saving {output_path}: {e}")

                # Add image to Word document
                try:
                    doc.add_picture(output_path, width=Inches(7.5))  # Adjust width to fit A4 page
                    doc.add_paragraph(f"Image {index}")
                    doc.add_page_break()
                except Exception as e:
                    print(f"Error adding {output_path} to document: {e}")

            # Update progress bar
            update_progress_bar(index, total_images)

        # Remove the last page break
        if doc.paragraphs[-1].text == '':
            doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)

        # Save the Word document
        word_output_path = os.path.join(input_folder, 'enhanced_images.docx')
        try:
            doc.save(word_output_path)
        except Exception as e:
            print(f"Error saving Word document: {e}")

        # Convert the Word document to PDF
        pdf_output_path = os.path.join(input_folder, 'enhanced_images.pdf')
        try:
            convert(word_output_path, pdf_output_path)
        except Exception as e:
            print(f"Error converting to PDF: {e}")

        # Open the PDF file
        try:
            os.startfile(pdf_output_path)
        except Exception as e:
            print(f"Error opening PDF file: {e}")

        messagebox.showinfo("Success", "Processing completed successfully.")
    
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")

# Create main tkinter window
root = Tk()
root.title("pic2doc-Grobern Business Solutions")

# Create a label and button in the window
label = Label(root, text="Select the folder containing images:")
label.pack(pady=10)

button = Button(root, text="Browse", command=select_folder)
button.pack(pady=10)

# Create a progress bar widget
progress_bar = ttk.Progressbar(root, orient='horizontal', length=300, mode='determinate')
progress_bar.pack(pady=10)

# Run the tkinter main loop
root.mainloop()
